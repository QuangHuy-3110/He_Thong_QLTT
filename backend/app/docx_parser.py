import os
import re
import docx

def parse_docx_lesson_plan(file_path):
    """
    Parses a lesson plan Word document (.docx) and extracts core metadata.
    Matches standard format from doc templates (Tên chủ đề, Môn, Lớp, Mô tả tóm tắt, Đối tượng, Loại hình).
    Returns clean metadata and attributes dict for automated form filling.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at {file_path}")

    doc = docx.Document(file_path)
    
    # Extract all paragraph texts cleanly
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs).lower()
    
    # Extract table texts
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt and txt not in table_texts:
                    table_texts.append(txt)
    full_table_text = "\n".join(table_texts).lower()
    combined_text = (full_text + "\n" + full_table_text)

    # Detect lesson plan document type
    lesson_plan_keywords = [
        "giáo án", "kế hoạch bài dạy", "kế hoạch dạy học", "mục tiêu dạy học", 
        "thiết bị dạy học", "học liệu", "tiến trình dạy học", "tiến trình dạy",
        "hoạt động 1", "hoạt động 01", "yêu cầu cần đạt", "sản phẩm trải nghiệm"
    ]
    is_lesson_plan = any(kw in combined_text for kw in lesson_plan_keywords)

    # 1. Extract Clean Title (Tên chủ đề / Bài)
    title = ""
    for p in paragraphs[:5]:
        p_clean = p.strip()
        if re.search(r"^(tên\s+chủ\s+đề|chủ\s+đề|bài\s+học|bài)\s*:", p_clean, re.IGNORECASE):
            title = re.sub(r"^(tên\s+chủ\s+đề|chủ\s+đề|bài\s+học|bài)\s*:\s*", "", p_clean, flags=re.IGNORECASE).strip()
            break
    if not title and paragraphs:
        title = paragraphs[0]
        title = re.sub(r"^(tên\s+chủ\s+đề|chủ\s+đề|bài\s+học|bài)\s*:\s*", "", title, flags=re.IGNORECASE).strip()

    # 2. Extract Subject (Môn học)
    subject = ""
    for p in paragraphs[:15]:
        # Match "Môn: Hoạt động trải nghiệm hướng nghiệp; lớp: 10" or "Môn học: Sinh học"
        if re.search(r"(?:môn\s*học|môn)\s*:", p, re.IGNORECASE):
            raw_sub = re.sub(r"^.*?(?:môn\s*học|môn)\s*:\s*", "", p, flags=re.IGNORECASE).strip()
            # Split by semicolon if present (e.g. "Hoạt động trải nghiệm hướng nghiệp; lớp: 10")
            parts = raw_sub.split(';')
            clean_first = parts[0].strip()
            # Remove any trailing "lớp: ..." if separated by space
            clean_first = re.sub(r"\s+lớp\s*:.*$", "", clean_first, flags=re.IGNORECASE).strip()
            if clean_first and len(clean_first) > 1:
                subject = clean_first
                break
    if not subject:
        if "sinh học" in combined_text and ("trải nghiệm" in combined_text or "hướng nghiệp" in combined_text):
            subject = "Hoạt động trải nghiệm Sinh học"
        elif "sinh học" in combined_text:
            subject = "Sinh học"
        elif "khoa học tự nhiên" in combined_text:
            subject = "Khoa học tự nhiên"
        elif "trải nghiệm" in combined_text:
            subject = "Hoạt động trải nghiệm, hướng nghiệp"
        else:
            subject = "Hoạt động trải nghiệm Sinh học"

    # 3. Extract Grade/Class (Cấp lớp) and normalize (e.g. "10" -> "Lớp 10")
    grade = ""
    for p in paragraphs[:10]:
        match = re.search(r"lớp\s*:\s*([^;,\n]+)", p, re.IGNORECASE)
        if match:
            raw_g = match.group(1).strip()
            num_match = re.search(r"\d+", raw_g)
            if num_match:
                grade = f"Lớp {num_match.group(0)}"
            else:
                grade = raw_g if raw_g.lower().startswith("lớp") else f"Lớp {raw_g}"
            break
    if not grade:
        grade = "Lớp 10"

    # 4. Extract Duration (Thời gian thực hiện)
    duration = ""
    for p in paragraphs[:10]:
        match = re.search(r"thời gian thực hiện[^:]*:\s*([^;\n]+)", p, re.IGNORECASE)
        if match:
            duration = match.group(1).strip()
            break

    # 5. Extract Summary / Description (Mô tả tóm tắt)
    description = ""
    desc_paragraphs = []
    capturing_desc = False

    for p in paragraphs:
        p_clean = p.strip()
        if re.search(r"^mô\s+tả\s+tóm\s+tắt\s*:", p_clean, re.IGNORECASE):
            capturing_desc = True
            text_after = re.sub(r"^mô\s+tả\s+tóm\s+tắt\s*:\s*", "", p_clean, flags=re.IGNORECASE).strip()
            if text_after:
                desc_paragraphs.append(text_after)
            continue
        
        if capturing_desc:
            # Stop capturing if we hit another metadata header or main section heading
            if re.search(r"^(đối\s+tượng|loại\s+hình|giao\s+viên|môn|thời\s+gian|[IVXLCDM]+\.|\d+\.)", p_clean, re.IGNORECASE):
                break
            desc_paragraphs.append(p_clean)

    if desc_paragraphs:
        description = " ".join(desc_paragraphs).strip()

    # Fallback description if not explicitly extracted
    if not description:
        if is_lesson_plan:
            desc_parts = []
            if subject: desc_parts.append(f"Bài giảng môn {subject}")
            if grade: desc_parts.append(f"dành cho {grade}")
            if duration: desc_parts.append(f"Thời gian thực hiện: {duration}")
            description = ", ".join(desc_parts) + "."
        else:
            desc_parts = [p for p in paragraphs[1:4] if len(p) > 20 and not p.startswith("Môn:")]
            description = " ".join(desc_parts)[:250] + "..." if desc_parts else "Tài liệu giáo án."

    # 6. Extract Target Students (Đối tượng giảng dạy)
    target_students = []
    for p in paragraphs[:15]:
        if re.search(r"đối\s+tượng\s+(giảng\s+dạy|học\s+sinh)\s*:", p, re.IGNORECASE):
            p_lower = p.lower()
            if "thành thị" in p_lower:
                target_students.append("Học sinh thành thị")
            if "nông thôn" in p_lower:
                target_students.append("Học sinh nông thôn")
            break

    if not target_students:
        urban_keywords = ["thành thị", "đô thị", "phố", "siêu thị", "trà sữa", "đồ ăn nhanh", "fast food", "sức khỏe học đường", "ít vận động"]
        rural_keywords = ["nông thôn", "làng", "bản", "ruộng", "vườn", "nông nghiệp"]
        if any(kw in combined_text for kw in urban_keywords):
            target_students.append("Học sinh thành thị")
        if any(kw in combined_text for kw in rural_keywords):
            target_students.append("Học sinh nông thôn")
            
    if not target_students:
        target_students = ["Học sinh thành thị", "Học sinh nông thôn"]

    # 7. Extract Lesson Type (Loại hình tiết dạy)
    lesson_type = ""
    for p in paragraphs[:15]:
        match = re.search(r"loại\s+hình\s+(tiết\s+dạy|bài\s+học|giảng\s+dạy)\s*:\s*(.+)", p, re.IGNORECASE)
        if match:
            lesson_type = match.group(2).strip()
            break

    if not lesson_type:
        if "lý thuyết" in combined_text:
            lesson_type = "Lý thuyết"
        elif "ôn tập" in combined_text:
            lesson_type = "Ôn tập"
        elif "kiểm tra" in combined_text:
            lesson_type = "Kiểm tra"
        elif "thực hành" in combined_text or "trải nghiệm" in combined_text:
            lesson_type = "Thực hành"
        else:
            lesson_type = "Hoạt động giáo dục theo chủ đề"

    # 8. Extract Activities (Tiến trình dạy học)
    activities = []
    for table in doc.tables:
        if len(table.rows) > 1:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            is_timeline_table = any("hoạt động" in h or "hđ" in h or "thời gian" in h for h in headers)
            if is_timeline_table:
                for row in table.rows[1:10]:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2:
                        raw_name = cells[0]
                        act_desc = cells[1] if len(cells) > 1 else ""
                        
                        time_match = re.search(r"(\d+\s*phút)", raw_name + " " + act_desc, re.IGNORECASE)
                        act_time = time_match.group(1) if time_match else "15 phút"
                        
                        # Strip time info
                        clean_col1 = re.sub(r"\(\s*\d+\s*phút\s*\)", "", raw_name, flags=re.IGNORECASE).strip()
                        clean_col1 = re.sub(r"\d+\s*phút", "", clean_col1, flags=re.IGNORECASE).strip()
                        clean_col1 = re.sub(r"[\s\-:]+$", "", clean_col1).strip()

                        # Check if col1 is just activity header (e.g. "HĐ 1", "Hoạt động 1", "1")
                        is_just_hd_header = bool(re.match(r"^(hoạt\s*động|hđ)?\s*\d+$", clean_col1, re.IGNORECASE))
                        
                        if is_just_hd_header and act_desc:
                            # col1 is just "HĐ 1", real title is the entire text of col2!
                            desc_clean = " ".join([l.strip() for l in act_desc.split("\n") if l.strip()])
                            act_name = desc_clean
                            act_desc_clean = ""
                        else:
                            # Remove any leading "Hoạt động 1:", "HĐ 1" prefix from act_name
                            act_name = re.sub(r"^(hoạt\s*động|hđ)\s*\d+[\s:\-]*", "", clean_col1, flags=re.IGNORECASE).strip()
                            if not act_name:
                                act_name = clean_col1
                            act_desc_clean = act_desc.strip() if act_desc else "Hoạt động dạy học chi tiết."

                        if len(act_desc_clean) > 250:
                            sentences = re.split(r'(?<=[.!?])\s+', act_desc_clean)
                            act_desc_clean = " ".join(sentences[:2]).strip()
                            
                        activities.append({
                            "ten_hoat_dong": act_name,
                            "thoi_gian": act_time,
                            "tom_tat": act_desc_clean
                        })
                break

    if not activities:
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            match = re.match(r"^(Hoạt động\s+\d+|HĐ\s*\d+)\s*:\s*(.*)", text, re.IGNORECASE)
            if match:
                raw_title = match.group(2).strip()
                act_name = re.sub(r"^(hoạt\s*động|hđ)\s*\d+[\s:\-]*", "", raw_title, flags=re.IGNORECASE).strip()
                if not act_name:
                    act_name = raw_title
                act_time = "10 phút"
                act_desc = ""
                
                for j in range(1, 6):
                    if i + j < len(doc.paragraphs):
                        next_text = doc.paragraphs[i + j].text.strip()
                        if not next_text:
                            continue
                        time_match = re.search(r"(\d+\s*phút)", next_text, re.IGNORECASE)
                        if time_match:
                            act_time = time_match.group(1).strip()
                        if not act_desc and len(next_text) > 20 and "Mục tiêu" not in next_text and "Tổ chức" not in next_text:
                            act_desc = next_text
                
                if act_desc:
                    sentences = re.split(r'(?<=[.!?])\s+', act_desc)
                    act_desc = " ".join(sentences[:2]).strip()
                else:
                    act_desc = "Tổ chức hoạt động giảng dạy trải nghiệm thực tế."
                
                activities.append({
                    "ten_hoat_dong": act_name,
                    "thoi_gian": act_time,
                    "tom_tat": act_desc
                })
                if len(activities) >= 5:
                    break

    # 9. Extract Knowledge Tags
    knowledge_tags = []
    common_tags = [
        "Dinh dưỡng học đường", "Thực đơn khỏe mạnh", "Nhóm chất dinh dưỡng",
        "Hoạt động trải nghiệm", "Sức khỏe học đường", "Chế độ ăn uống",
        "Vận động", "Thói quen ăn uống", "Thiết kế thực đơn", "Trò chơi trải nghiệm",
        "Chăm sóc sức khỏe", "Thời gian biểu", "Đồng hồ sinh học", "Giấc ngủ", "Quản lý cảm xúc"
    ]
    for tag in common_tags:
        if tag.lower() in combined_text:
            knowledge_tags.append(tag)
    knowledge_tags = knowledge_tags[:6]

    # 10. Infer Experiential Curriculum Attributes (Mạch kiến thức)
    track = ""
    for p in paragraphs[:20]:
        match = re.search(r"mạch\s*kiến\s*thức\s*:\s*([^;\n]+)", p, re.IGNORECASE)
        if match:
            track = match.group(1).strip()
            break
    if not track:
        if any(k in combined_text for k in ["bản thân", "sức khỏe", "thể chất", "cơ thể", "giấc ngủ", "thói quen"]):
            track = "Hoạt động hướng vào bản thân"
        elif any(k in combined_text for k in ["gia đình", "nhà trường", "cộng đồng", "xã hội", "bè bạn", "giao tiếp"]):
            track = "Hoạt động hướng đến xã hội"
        elif any(k in combined_text for k in ["thiên nhiên", "môi trường", "cảnh quan", "động vật", "thực vật"]):
            track = "Hoạt động hướng đến tự nhiên"
        elif any(k in combined_text for k in ["nghề nghiệp", "định hướng nghề", "lao động", "nghề"]):
            track = "Hoạt động hướng nghiệp"
        else:
            track = "Hoạt động hướng vào bản thân"

    topic = title or "Khám phá bản thân"

    biology_connections = []
    bio_integration_details = []

    # Direct extraction from docx tables (KIẾN THỨC SINH HỌC ĐƯỢC TÍCH HỢP)
    for table in doc.tables:
        if len(table.rows) > 1:
            headers = [c.text.strip().lower() for c in table.rows[0].cells]
            if any(h == "kiến thức sh" or "kiến thức sh" in h or ("kiến thức sinh học" in h and "tích hợp" in h) for h in headers):
                bio_col_idx = 0
                for i, h in enumerate(headers):
                    if h == "kiến thức sh" or "kiến thức sh" in h or ("kiến thức sinh học" in h and "tích hợp" in h):
                        bio_col_idx = i
                        break
                for row in table.rows[1:]:
                    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
                    if len(cells) > bio_col_idx and cells[bio_col_idx]:
                        val = re.sub(r"^\d+\.\s*", "", cells[bio_col_idx]).strip()
                        if val and val not in biology_connections and len(val) > 2:
                            biology_connections.append(val)
                    
                    if len(cells) >= 2 and cells[0]:
                        bio_integration_details.append({
                            "kien_thuc_sh": cells[0],
                            "noi_dung_tich_hop": cells[1] if len(cells) > 1 else "",
                            "y_nghia": cells[2] if len(cells) > 2 else ""
                        })

    location = "Lớp học tiêu chuẩn"
    if "thí nghiệm" in combined_text:
        location = "Phòng thí nghiệm Sinh học"
    elif "sân trường" in combined_text or "ngoài trời" in combined_text:
        location = "Ngoài trời / Sân trường"
    elif "đa năng" in combined_text or "nhà ăn" in combined_text:
        location = "Phòng đa năng / Nhà ăn"

    attributes = {
        "Mạch kiến thức": track,
        "Chủ đề": topic,
        "Kiến thức sinh học liên quan": biology_connections,
        "bio_integration_details": bio_integration_details,
        "Địa điểm": location,
        "lop": [grade],
        "Loại hình": lesson_type,
        "Thời gian thực hiện": duration
    }

    return {
        "title": title,
        "description": description,
        "subject": subject,
        "grade": grade,
        "duration": duration,
        "target_students": target_students,
        "lesson_type": lesson_type,
        "activities": activities,
        "knowledge_tags": knowledge_tags,
        "attributes": attributes
    }

def convert_docx_to_markdown(file_path):
    """
    Reads a Word Document (.docx) sequentially and outputs its content as beautifully
    formatted Markdown, preserving lists, styles, and full tabular structures.
    """
    import docx
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    import re

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at {file_path}")

    doc = docx.Document(file_path)
    md_parts = []
    
    # Iterate through docx body elements sequentially
    for child in doc.element.body:
        if child.tag.endswith('p'):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if not text:
                continue
                
            # Formulate headings
            if text.upper().startswith("CHỦ ĐỀ") or text.upper().startswith("BÀI HỌC:"):
                md_parts.append(f"\n# {text}\n")
            elif re.match(r"^[IVXLCDM]+\.\s+", text):
                md_parts.append(f"\n## {text}\n")
            elif re.match(r"^(Hoạt động\s+\d+|HĐ\s*\d+|[0-9]+\.\s+)", text, re.IGNORECASE):
                md_parts.append(f"\n### {text}\n")
            elif text.startswith("-") or text.startswith("*") or text.startswith("•"):
                # Normalize list item
                clean_item = re.sub(r"^[-*•]\s*", "", text)
                md_parts.append(f"- {clean_item}")
            else:
                md_parts.append(text + "\n")
                
        elif child.tag.endswith('tbl'):
            t = Table(child, doc)
            if not t.rows:
                continue
                
            table_md = []
            
            # Formulate Table Header
            header_cells = t.rows[0].cells
            header_texts = []
            seen_headers = []
            for c in header_cells:
                ct = c.text.strip().replace('\n', ' ')
                # Avoid duplicate cell references in merged cells
                if not seen_headers or seen_headers[-1] != ct or ct == "":
                    header_texts.append(ct)
                    seen_headers.append(ct)
            
            # If all empty, skip
            if not any(header_texts):
                header_texts = [f"Cột {i+1}" for i in range(len(header_cells))]

            table_md.append("| " + " | ".join(header_texts) + " |")
            table_md.append("| " + " | ".join("---" for _ in header_texts) + " |")
            
            # Formulate Table Rows
            for row in t.rows[1:]:
                row_cells = row.cells
                row_texts = []
                seen_row = []
                for c in row_cells:
                    ct = c.text.strip().replace('\n', ' ')
                    # Avoid duplications from merged cells
                    if not seen_row or seen_row[-1] != ct or ct == "":
                        row_texts.append(ct)
                        seen_row.append(ct)
                # Pad row_texts if it is shorter than header
                while len(row_texts) < len(header_texts):
                    row_texts.append("")
                # Truncate if longer
                row_texts = row_texts[:len(header_texts)]
                
                table_md.append("| " + " | ".join(row_texts) + " |")
                
            md_parts.append("\n" + "\n".join(table_md) + "\n")

    return "\n".join(md_parts)

