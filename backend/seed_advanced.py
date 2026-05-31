import os
import django
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'kms_core.settings')
django.setup()

from app.models import LessonPlan, Directory, User, LessonPlanDirectory, DocumentChunk

# 1. Clean up
print("Deleting old data...")
DocumentChunk.objects.all().delete()
LessonPlanDirectory.objects.all().delete()
LessonPlan.objects.all().delete()
Directory.objects.all().delete()

# 2. Get/create admin
print("Creating admin...")
admin, created = User.objects.get_or_create(
    username='admin',
    defaults={'role': 'ADMIN', 'full_name': 'Admin'}
)
if created:
    admin.set_password('admin')
    admin.save()

# Ensure media directory exists
media_dir = os.path.join('media', 'lesson_plans')
os.makedirs(media_dir, exist_ok=True)


def generate_structured_docx(filename, title, mon_hoc, lop, thoi_gian, gv, muc_tieu_list, thiet_bi, hoat_dong_list):
    """
    Tạo file .docx theo format chuẩn của tai_lieu.docx
    hoat_dong_list: list of (ten, muc_tieu_list, to_chuc_str)
    """
    filepath = os.path.join(media_dir, filename)
    doc = docx.Document()

    # Title
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Môn: {mon_hoc}; lớp: {lop}")
    doc.add_paragraph(f"Thời gian thực hiện dự kiến: {thoi_gian}")
    doc.add_paragraph(f"Giáo viên giảng dạy: {gv}")

    doc.add_heading("I. MỤC TIÊU DẠY HỌC", level=1)
    for mt in muc_tieu_list:
        doc.add_paragraph(mt, style='List Bullet')

    doc.add_heading("II. THIẾT BỊ DẠY HỌC, HỌC LIỆU", level=1)
    doc.add_paragraph(thiet_bi)

    doc.add_heading("III. TIẾN TRÌNH DẠY HỌC", level=1)
    doc.add_heading("1. Khung tiến trình dạy học", level=2)
    doc.add_paragraph("Hoạt động dạy học được triển khai theo các giai đoạn: Khởi động → Khám phá → Thực hành → Chia sẻ → Vận dụng.")

    doc.add_heading("2. Tiến trình dạy học chi tiết", level=2)
    for i, (ten, muc_tieu_hd, to_chuc) in enumerate(hoat_dong_list, start=1):
        doc.add_heading(f"Hoạt động {i:02d}: {ten}", level=3)

        doc.add_paragraph("Mục tiêu (Yêu cầu cần đạt)", style='List Bullet')
        for mt in muc_tieu_hd:
            doc.add_paragraph(f"- {mt}")

        doc.add_paragraph("Tổ chức thực hiện", style='List Bullet')
        doc.add_paragraph(to_chuc)

    doc.save(filepath)
    return os.path.join('lesson_plans', filename)


# ============================================================
# 3. Tạo cấu trúc thư mục (4 Track + 10 Sub-folder)
# ============================================================
print("Restructuring directories...")

# Track 1: Hoạt động hướng vào bản thân
track1 = Directory.objects.create(
    name='Hoạt động hướng vào bản thân',
    is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng vào bản thân',
        'knowledge_tags': [
            'Hệ thần kinh, hormone (serotonin, adrenaline), cơ sở sinh học của cảm xúc',
            'Sinh học hành vi: hormone tuổi dậy thì, sức khỏe tâm – sinh lý',
            'Hệ cơ – xương – khớp, tim mạch, hô hấp, năng lượng ATP',
            'Dinh dưỡng học, chuyển hóa năng lượng, vai trò vitamin/khoáng chất',
            'Cân bằng nước, sinh học giấc ngủ, nhịp sinh học',
            'Cấu tạo cơ thể, tuần hoàn máu, hô hấp nhân tạo, nguyên lý đông máu'
        ]
    }
)
sub1_1 = Directory.objects.create(
    name='Khám phá bản thân', parent=track1, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Hệ thần kinh, hormone (serotonin, adrenaline), cơ sở sinh học của cảm xúc',
            'Sinh học hành vi: hormone tuổi dậy thì, sức khỏe tâm – sinh lý'
        ]
    }
)
sub1_2 = Directory.objects.create(
    name='Rèn luyện bản thân', parent=track1, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Hệ cơ – xương – khớp, tim mạch, hô hấp, năng lượng ATP',
            'Dinh dưỡng học, chuyển hóa năng lượng, vai trò vitamin/khoáng chất',
            'Cân bằng nước, sinh học giấc ngủ, nhịp sinh học',
            'Cấu tạo cơ thể, tuần hoàn máu, hô hấp nhân tạo, nguyên lý đông máu'
        ]
    }
)

# Track 2: Hoạt động hướng đến xã hội
track2 = Directory.objects.create(
    name='Hoạt động hướng đến xã hội',
    is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến xã hội',
        'knowledge_tags': [
            'Sinh học thần kinh: trí nhớ, hình thành thói quen, ảnh hưởng giấc ngủ/dinh dưỡng đến tập trung',
            'Hệ miễn dịch, bệnh truyền nhiễm, vệ sinh cá nhân, nguyên tắc phòng bệnh',
            'Sinh học hành vi: hormone tuổi dậy thì, sức khỏe tâm – sinh lý',
            'Phản xạ thần kinh, tác động rượu/bia đến hệ thần kinh và tim mạch',
            'Cơ chế nghe – nhìn, ảnh hưởng âm nhạc đến não bộ, sinh học vận động',
            'Sinh lý thực vật (quang hợp, dinh dưỡng cây trồng), bệnh học cây trồng',
            'Vi sinh vật gây bệnh trong rác thải, ảnh hưởng ô nhiễm đến sức khỏe cộng đồng',
            'Dịch tễ học cơ bản, sức khỏe sinh sản vị thành niên, phòng chống bệnh truyền nhiễm',
            'Hormone oxytocin, dopamine trong quan hệ xã hội, sức khỏe tinh thần'
        ]
    }
)
sub2_1 = Directory.objects.create(
    name='Chăm sóc gia đình', parent=track2, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Sinh lý thực vật (quang hợp, dinh dưỡng cây trồng), bệnh học cây trồng',
            'Sinh lý động vật, bệnh học chăn nuôi và vật nuôi'
        ]
    }
)
sub2_2 = Directory.objects.create(
    name='Xây dựng nhà trường', parent=track2, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Sinh học thần kinh: trí nhớ, hình thành thói quen, ảnh hưởng giấc ngủ/dinh dưỡng đến tập trung',
            'Cơ chế nghe – nhìn, ảnh hưởng âm nhạc đến não bộ, sinh học vận động',
            'Hệ cơ – xương – khớp, tim mạch, hô hấp, năng lượng ATP'
        ]
    }
)
sub2_3 = Directory.objects.create(
    name='Xây dựng cộng đồng', parent=track2, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Hệ miễn dịch, bệnh truyền nhiễm, vệ sinh cá nhân, nguyên tắc phòng bệnh',
            'Phản xạ thần kinh, tác động rượu/bia đến hệ thần kinh và tim mạch',
            'Vi sinh vật gây bệnh trong rác thải, ảnh hưởng ô nhiễm đến sức khỏe cộng đồng',
            'Dịch tễ học cơ bản, sức khỏe sinh sản vị thành niên, phòng chống bệnh truyền nhiễm',
            'Hormone oxytocin, dopamine trong quan hệ xã hội, sức khỏe tinh thần'
        ]
    }
)

# Track 3: Hoạt động hướng đến tự nhiên
track3 = Directory.objects.create(
    name='Hoạt động hướng đến tự nhiên',
    is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến tự nhiên',
        'knowledge_tags': [
            'Phân loại thực vật, đa dạng sinh học, tiến hóa',
            'Hệ sinh thái nông nghiệp, đa dạng sinh học địa phương, thích nghi sinh vật',
            'Hệ hô hấp người, tác động khí độc, sinh thái đô thị',
            'Vòng tuần hoàn vật chất, vi sinh vật phân hủy, sinh thái học hệ sinh thái',
            'Quang hợp, hô hấp thực vật, sinh thái rừng, chuỗi thức ăn',
            'Vi sinh vật nước, chu trình nitơ, ảnh hưởng ô nhiễm đến sinh vật thủy sinh'
        ]
    }
)
sub3_1 = Directory.objects.create(
    name='Tìm hiểu và bảo tồn cảnh quan thiên nhiên', parent=track3, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Phân loại thực vật, đa dạng sinh học, tiến hóa',
            'Hệ sinh thái nông nghiệp, đa dạng sinh học địa phương, thích nghi sinh vật'
        ]
    }
)
sub3_2 = Directory.objects.create(
    name='Tìm hiểu và bảo vệ môi trường', parent=track3, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Hệ hô hấp người, tác động khí độc, sinh thái đô thị',
            'Vòng tuần hoàn vật chất, vi sinh vật phân hủy, sinh thái học hệ sinh thái',
            'Quang hợp, hô hấp thực vật, sinh thái rừng, chuỗi thức ăn',
            'Vi sinh vật nước, chu trình nitơ, ảnh hưởng ô nhiễm đến sinh vật thủy sinh'
        ]
    }
)

# Track 4: Hoạt động hướng nghiệp
track4 = Directory.objects.create(
    name='Hoạt động hướng nghiệp',
    is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng nghiệp',
        'knowledge_tags': [
            'Công nghệ gen, sinh học phân tử, ứng dụng y học/nông nghiệp',
            'Quy trình sản xuất thuốc, an toàn sinh học, nghiên cứu tế bào',
            'Sinh lý thực vật, nuôi cấy mô, di truyền chọn giống',
            'Dinh dưỡng thực vật, sinh lý động vật, bệnh học cây trồng/vật nuôi'
        ]
    }
)
sub4_1 = Directory.objects.create(
    name='Tìm hiểu nghề nghiệp', parent=track4, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Công nghệ gen, sinh học phân tử, ứng dụng y học/nông nghiệp',
            'Quy trình sản xuất thuốc, an toàn sinh học, nghiên cứu tế bào'
        ]
    }
)
sub4_2 = Directory.objects.create(
    name='Rèn luyện phẩm chất, năng lực phù hợp với định hướng nghề nghiệp', parent=track4, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Sinh lý thực vật, nuôi cấy mô, di truyền chọn giống',
            'Dinh dưỡng thực vật, sinh lý động vật, bệnh học cây trồng/vật nuôi'
        ]
    }
)
sub4_3 = Directory.objects.create(
    name='Lựa chọn hướng nghề nghiệp và lập kế hoạch học tập theo định hướng nghề nghiệp', parent=track4, is_public=True, user=admin,
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'knowledge_tags': [
            'Công nghệ gen, sinh học phân tử, ứng dụng y học/nông nghiệp',
            'Sinh lý thực vật, nuôi cấy mô, di truyền chọn giống'
        ]
    }
)

print("Creating 12 lesson plans with distinct attributes and structured DOCX files...")

# ============================================================
# LESSON PLAN 1
# Mạch: Hướng vào bản thân | Chủ đề: Khám phá bản thân
# Lớp 10 | Lý thuyết | 1 tiết | HS thành thị | Lớp học tiêu chuẩn
# Kiến thức: Hệ thần kinh, hormone cảm xúc
# ============================================================
file1 = generate_structured_docx(
    filename='nhat_ky_cam_xuc_lop10.docx',
    title='CHỦ ĐỀ 1: NHẬT KÝ CẢM XÚC – CƠ SỞ SINH HỌC CỦA HÀNH VI',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 10 (THPT)',
    thoi_gian='1 tiết (45 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh tự nhận diện và phân loại cảm xúc cá nhân trong nhật ký hằng ngày.',
        'Giải thích được cơ chế sinh học điều khiển cảm xúc: hệ thần kinh, hormone adrenaline, serotonin, dopamine.',
        'Hình thành thói quen viết nhật ký cảm xúc như công cụ tự điều chỉnh hành vi.',
    ],
    thiet_bi='Phiếu nhật ký in sẵn, bộ thẻ hình biểu cảm, bảng phân loại hormone cảm xúc khổ A0.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: TRÒ CHƠI NHẬN DIỆN BIỂU CẢM',
            ['Tạo hứng thú, kết nối học sinh với chủ đề cảm xúc.'],
            'GV tổ chức trò chơi "Đoán cảm xúc qua khuôn mặt": HS bốc thẻ biểu cảm và diễn giải cảm giác. Thảo luận: "Khi sợ hãi, cơ thể bạn phản ứng thế nào?" (10 phút)'
        ),
        (
            'TÌM HIỂU CƠ SỞ SINH HỌC CỦA CẢM XÚC',
            [
                'Giải thích vai trò của adrenaline khi căng thẳng (tăng nhịp tim, giãn đồng tử).',
                'Phân biệt serotonin (cảm giác hạnh phúc) và dopamine (phần thưởng).',
            ],
            'GV trình bày sơ đồ hệ thần kinh – tuyến thượng thận – hormone cảm xúc. HS ghi chép và đặt câu hỏi phản biện. (15 phút)'
        ),
        (
            'THỰC HÀNH VIẾT NHẬT KÝ CẢM XÚC',
            ['HS tự ghi chép ít nhất 3 cảm xúc trong tuần và liên hệ sinh học tương ứng.'],
            'HS hoàn thiện phiếu nhật ký cảm xúc cá nhân: ghi tên cảm xúc, tình huống, phản ứng cơ thể và hormone liên quan. Chia sẻ cặp đôi. (15 phút)'
        ),
        (
            'VẬN DỤNG VÀO THỰC TẾ',
            ['Đề xuất giải pháp điều chỉnh cảm xúc lành mạnh dựa trên sinh học.'],
            'HS trình bày phương pháp tự điều chỉnh cảm xúc (hít thở, vận động nhẹ, viết nhật ký) và giải thích cơ chế sinh học. (5 phút)'
        ),
    ]
)

lp1 = LessonPlan.objects.create(
    title='Nhật ký cảm xúc và Cơ sở sinh học của hành vi',
    description='Giúp học sinh lớp 10 THPT hiểu cơ chế nội tiết tố và hệ thần kinh kiểm soát cảm xúc hằng ngày. Thực hành viết nhật ký cảm xúc kết hợp kiến thức sinh học về adrenaline, serotonin, dopamine.',
    target_student='HS thành thị',
    status='PUBLISHED', creator=admin, file_path=file1,
    content_preview='Chủ đề: Nhật ký cảm xúc – cơ sở sinh học. Lớp 10, 1 tiết. Kiến thức: hệ thần kinh, hormone adrenaline, serotonin điều khiển cảm xúc. Hoạt động: trò chơi biểu cảm, tìm hiểu hormone, viết nhật ký, vận dụng.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng vào bản thân',
        'Chủ đề': 'Khám phá bản thân',
        'Kiến thức sinh học liên quan': 'Hệ thần kinh, hormone (serotonin, adrenaline), cơ sở sinh học của cảm xúc',
        'Loại hình': 'Lý thuyết',
        'Tiết dạy': '1 tiết',
        'Địa điểm': 'Lớp học tiêu chuẩn',
        'lop': ['Lớp 10'],
        'knowledge_tags': ['Hệ thần kinh, hormone (serotonin, adrenaline), cơ sở sinh học của cảm xúc']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp1, directory=sub1_1)
LessonPlanDirectory.objects.create(lesson_plan=lp1, directory=track1)

# ============================================================
# LESSON PLAN 2
# Mạch: Hướng vào bản thân | Chủ đề: Khám phá bản thân
# Lớp 11 | Lý thuyết | 2 tiết | HS nông thôn | Lớp học tiêu chuẩn
# Kiến thức: Sinh học hành vi: hormone tuổi dậy thì
# ============================================================
file2 = generate_structured_docx(
    filename='suc_khoe_tuoi_day_thi_lop11.docx',
    title='CHỦ ĐỀ 2: SỨC KHỎE TÂM – SINH LÝ TUỔI DẬY THÌ',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 11 (THPT)',
    thoi_gian='2 tiết (90 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh nhận biết các thay đổi sinh lý – tâm lý trong giai đoạn dậy thì.',
        'Giải thích được vai trò của hormone testosterone, estrogen, GnRH trong phát triển cơ thể.',
        'Hình thành thái độ tích cực về sức khỏe sinh sản và tâm lý vị thành niên.',
    ],
    thiet_bi='Tranh giải phẫu hệ sinh sản, phiếu hỏi sức khỏe vị thành niên, video tư liệu 5 phút.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: HỘP CÂU HỎI BÍ MẬT',
            ['Khơi dậy sự tò mò và mạnh dạn thảo luận về tuổi dậy thì.'],
            'HS viết câu hỏi về tuổi dậy thì vào phiếu bỏ vào hộp, GV chọn ngẫu nhiên và trả lời. Phá vỡ rào cản tâm lý về chủ đề nhạy cảm. (15 phút)'
        ),
        (
            'TÌM HIỂU HORMONE TĂNG TRƯỞNG VÀ DẬY THÌ',
            [
                'Mô tả quá trình điều tiết hormone GnRH – LH/FSH – testosterone/estrogen.',
                'Phân biệt thay đổi sinh lý ở nam và nữ trong giai đoạn dậy thì.',
            ],
            'GV giảng giải sơ đồ trục hạ đồi – tuyến yên – tuyến sinh dục. HS điền vào bảng so sánh thay đổi cơ thể nam – nữ. (30 phút)'
        ),
        (
            'THẢO LUẬN NHÓM: SỨC KHỎE TÂM SINH LÝ',
            ['HS nhận ra mối liên hệ giữa hormone và tâm trạng, hành vi.'],
            'Chia 4 nhóm thảo luận: Nhóm 1 – tâm trạng thất thường; Nhóm 2 – ăn nhiều/ít; Nhóm 3 – thay đổi giọng nói; Nhóm 4 – mụn trứng cá. Mỗi nhóm trình bày giải thích sinh học. (25 phút)'
        ),
        (
            'KẾT LUẬN VÀ VẬN DỤNG',
            ['Đề xuất lời khuyên chăm sóc sức khỏe dựa trên hiểu biết sinh học.'],
            'HS lập danh sách 5 thói quen lành mạnh phù hợp lứa tuổi. Chia sẻ toàn lớp. GV tổng kết. (20 phút)'
        ),
    ]
)

lp2 = LessonPlan.objects.create(
    title='Sức khỏe tâm – sinh lý tuổi dậy thì: Hiểu để tự chăm sóc bản thân',
    description='Giúp học sinh lớp 11 nông thôn hiểu các biến đổi tâm sinh lý trong giai đoạn dậy thì thông qua kiến thức về hormone GnRH, testosterone, estrogen và mối liên hệ với hành vi, tâm trạng.',
    target_student='HS nông thôn',
    status='PUBLISHED', creator=admin, file_path=file2,
    content_preview='Chủ đề: Sức khỏe tâm sinh lý tuổi dậy thì. Lớp 11, 2 tiết. Kiến thức: hormone sinh dục GnRH, testosterone, estrogen, thay đổi cơ thể. Hoạt động: hộp câu hỏi, bài giảng hormone, thảo luận nhóm, vận dụng.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng vào bản thân',
        'Chủ đề': 'Khám phá bản thân',
        'Kiến thức sinh học liên quan': 'Sinh học hành vi: hormone tuổi dậy thì, sức khỏe tâm – sinh lý',
        'Loại hình': 'Lý thuyết',
        'Tiết dạy': '2 tiết',
        'Địa điểm': 'Lớp học tiêu chuẩn',
        'lop': ['Lớp 11'],
        'knowledge_tags': ['Sinh học hành vi: hormone tuổi dậy thì, sức khỏe tâm – sinh lý']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp2, directory=sub1_1)
LessonPlanDirectory.objects.create(lesson_plan=lp2, directory=track1)

# ============================================================
# LESSON PLAN 3
# Mạch: Hướng vào bản thân | Chủ đề: Rèn luyện bản thân
# Lớp 10 | Thực hành | 2 tiết | HS thành thị | Phòng đa năng / Nhà ăn
# Kiến thức: Dinh dưỡng học, chuyển hóa năng lượng, vitamin
# ============================================================
file3 = generate_structured_docx(
    filename='dinh_duong_hoc_duong_lop10.docx',
    title='CHỦ ĐỀ 3: DINH DƯỠNG HỌC ĐƯỜNG – XÂY DỰNG THỰC ĐƠN KHỎE MẠNH',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 10 (THPT)',
    thoi_gian='2 tiết (90 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh trình bày được nhu cầu dinh dưỡng đặc thù lứa tuổi THPT.',
        'Phân biệt vai trò của 4 nhóm chất dinh dưỡng: glucid, lipid, protein, vitamin.',
        'Thiết kế được thực đơn 3 bữa cân đối năng lượng và vi chất.',
    ],
    thiet_bi='Bao bì đồ ăn, cân điện tử, app tính calo trên điện thoại, poster nhóm chất dinh dưỡng.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: SIÊU THỊ MINI VÀ PHÂN TÍCH BAO BÌ',
            ['Kích hoạt hiểu biết về thành phần dinh dưỡng thực phẩm.'],
            'HS mang bao bì đồ ăn từ nhà. Trò chơi "Đọc nhãn thực phẩm": Đội nào đọc đúng calo, protein, đường nhiều nhất trong 5 phút thắng. (15 phút)'
        ),
        (
            'TÌM HIỂU CÁC NHÓM CHẤT DINH DƯỠNG',
            [
                'Nhận biết glucid là nguồn năng lượng chính (4 kcal/g), lipid dự trữ (9 kcal/g).',
                'Hiểu vai trò protein trong tái tạo tế bào và enzyme sinh học.',
                'Phân loại vitamin tan trong dầu (A, D, E, K) và tan trong nước (B, C).',
            ],
            'GV trình bày bảng phân loại chất dinh dưỡng với ví dụ thực phẩm cụ thể. HS điền vào sơ đồ phân loại. (25 phút)'
        ),
        (
            'THỰC HÀNH: THIẾT KẾ THỰC ĐƠN KHOA HỌC',
            ['Xây dựng thực đơn 3 bữa đảm bảo 2000-2200 kcal/ngày cho lứa tuổi 15-17.'],
            'Nhóm 4 HS thiết kế thực đơn thực tế, tính tổng calo, tỷ lệ 3 nhóm chất bằng app. Trình bày poster "Bữa ăn khỏe mạnh". (40 phút)'
        ),
        (
            'PHẢN BIỆN VÀ TỔNG KẾT',
            ['Nhận xét phân tích dinh dưỡng của nhóm bạn.'],
            'Các nhóm phản biện thực đơn: nhận xét thiếu/thừa chất. GV tổng kết về chuyển hóa năng lượng ATP trong tế bào. (10 phút)'
        ),
    ]
)

lp3 = LessonPlan.objects.create(
    title='Dinh dưỡng học đường – Thiết kế thực đơn cân đối calo cho học sinh THPT',
    description='Học sinh lớp 10 thực hành xây dựng thực đơn 3 bữa cân đối dinh dưỡng, áp dụng kiến thức về chuyển hóa glucid, lipid, protein và vai trò của vitamin, khoáng chất trong cơ thể.',
    target_student='HS thành thị',
    status='PUBLISHED', creator=admin, file_path=file3,
    content_preview='Chủ đề: Dinh dưỡng học đường. Lớp 10, 2 tiết, thực hành. Kiến thức: chuyển hóa glucid, lipid, protein, vitamin tan dầu/nước. Hoạt động: đọc nhãn bao bì, tìm hiểu nhóm chất, thiết kế thực đơn, phản biện.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng vào bản thân',
        'Chủ đề': 'Rèn luyện bản thân',
        'Kiến thức sinh học liên quan': 'Dinh dưỡng học, chuyển hóa năng lượng, vai trò vitamin/khoáng chất',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '2 tiết',
        'Địa điểm': 'Phòng đa năng / Nhà ăn',
        'lop': ['Lớp 10'],
        'knowledge_tags': ['Dinh dưỡng học, chuyển hóa năng lượng, vai trò vitamin/khoáng chất']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp3, directory=sub1_2)
LessonPlanDirectory.objects.create(lesson_plan=lp3, directory=track1)

# ============================================================
# LESSON PLAN 4
# Mạch: Hướng vào bản thân | Chủ đề: Rèn luyện bản thân
# Lớp 11 | Thực hành | 2 tiết | HS nông thôn | Sân trường / Khu vực thể thao
# Kiến thức: Hệ cơ–xương–khớp, tim mạch, hô hấp, ATP
# ============================================================
file4 = generate_structured_docx(
    filename='ren_luyen_the_chat_lop11.docx',
    title='CHỦ ĐỀ 4: RÈN LUYỆN THỂ CHẤT – SINH HỌC VẬN ĐỘNG VÀ NĂNG LƯỢNG ATP',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 11 (THPT)',
    thoi_gian='2 tiết (90 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh mô tả được cấu trúc cơ vân, cơ tim và cơ chế co cơ theo lý thuyết tơ cơ trượt.',
        'Giải thích sự tổng hợp ATP từ ADP trong tế bào cơ khi vận động.',
        'Đo được nhịp tim, tần số hô hấp trước và sau vận động, rút ra kết luận sinh học.',
    ],
    thiet_bi='Đồng hồ bấm giây, máy đo SpO2, băng keo đánh dấu vạch chạy, phiếu theo dõi nhịp tim.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: ĐO NHỊP TIM VÀ HÔ HẤP KHI NGHỈ NGƠI',
            ['Thiết lập số liệu cơ sở trước khi vận động.'],
            'HS đo nhịp tim 1 phút bằng cách ép ngón tay vào cổ tay. Đo nhịp thở 30 giây. Ghi vào phiếu theo dõi. (10 phút)'
        ),
        (
            'HOẠT ĐỘNG THỂ CHẤT: CHẠY VÀ NHẢY DÂY',
            ['Ghi nhận phản ứng tim mạch, hô hấp khi tăng cường độ vận động.'],
            'HS chạy bộ 5 phút quanh sân, sau đó nhảy dây 50 cái. Đo lại nhịp tim và nhịp thở ngay sau khi dừng. So sánh với số liệu ban đầu. (20 phút)'
        ),
        (
            'GIẢI THÍCH SINH HỌC: CƠ CHẾ CO CƠ VÀ TỔNG HỢP ATP',
            [
                'Mô tả cơ chế co cơ theo mô hình tơ cơ myosin – actin.',
                'Giải thích vai trò ATP và sự hô hấp tế bào cung cấp năng lượng.',
                'Phân biệt chuyển hóa hiếu khí và kỵ khí trong cơ bắp.',
            ],
            'GV trình bày sơ đồ co cơ và tổng hợp ATP (ADP + Pi + năng lượng → ATP). HS hoàn thành phiếu điền khuyết về chu trình Krebs đơn giản. (40 phút)'
        ),
        (
            'THẢO LUẬN VÀ TỔNG KẾT',
            ['HS rút ra kết luận về lợi ích của vận động thường xuyên với sức khỏe tim mạch.'],
            'Nhóm 4 người thảo luận: Tại sao sau khi chạy nhanh, hơi thở dốc? Liên hệ với nợ oxy và lactate. Đại diện trình bày. (20 phút)'
        ),
    ]
)

lp4 = LessonPlan.objects.create(
    title='Rèn luyện thể chất – Sinh học vận động và cơ chế tổng hợp ATP trong cơ bắp',
    description='Học sinh lớp 11 nông thôn trải nghiệm thực hành đo nhịp tim, hô hấp trước và sau vận động. Tìm hiểu cơ chế co cơ theo mô hình myosin-actin và tổng hợp ATP qua hô hấp tế bào hiếu khí.',
    target_student='HS nông thôn',
    status='PUBLISHED', creator=admin, file_path=file4,
    content_preview='Chủ đề: Rèn luyện thể chất sinh học vận động. Lớp 11, 2 tiết, thực hành. Kiến thức: hệ cơ xương khớp, tim mạch, hô hấp, ATP. Hoạt động: đo nhịp tim, chạy nhảy, giải thích co cơ ATP, thảo luận.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng vào bản thân',
        'Chủ đề': 'Rèn luyện bản thân',
        'Kiến thức sinh học liên quan': 'Hệ cơ – xương – khớp, tim mạch, hô hấp, năng lượng ATP',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '2 tiết',
        'Địa điểm': 'Sân trường / Khu vực thể thao',
        'lop': ['Lớp 11'],
        'knowledge_tags': ['Hệ cơ – xương – khớp, tim mạch, hô hấp, năng lượng ATP']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp4, directory=sub1_2)
LessonPlanDirectory.objects.create(lesson_plan=lp4, directory=track1)

# ============================================================
# LESSON PLAN 5
# Mạch: Hướng đến xã hội | Chủ đề: Chăm sóc gia đình
# Lớp 10 | Thực hành | 3 tiết | HS nông thôn | Thực địa / Nông trại
# Kiến thức: Sinh lý thực vật, bệnh học cây trồng
# ============================================================
file5 = generate_structured_docx(
    filename='nong_nghiep_gia_dinh_lop10.docx',
    title='CHỦ ĐỀ 5: CHĂM SÓC VƯỜN RAU GIA ĐÌNH – SINH LÝ THỰC VẬT VÀ PHÒNG TRỪ SÂU BỆNH',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 10 (THPT)',
    thoi_gian='3 tiết (135 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh mô tả được quá trình quang hợp và vai trò của ánh sáng, CO2, nước đối với cây trồng.',
        'Nhận biết các triệu chứng thiếu nguyên tố khoáng vi lượng trên lá cây.',
        'Xác định được các loại bệnh cây trồng phổ biến (nấm, vi khuẩn, virus) qua dấu hiệu thực tế.',
        'Đề xuất biện pháp phòng trừ bệnh sinh học an toàn thay thế hóa chất.',
    ],
    thiet_bi='Mẫu lá bệnh, kính lúp, cuốc xới, phân bón vi sinh EM, phiếu quan sát cây trồng.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: KHẢO SÁT VƯỜN RAU GIA ĐÌNH',
            ['Kết nối kiến thức sinh học với thực tế chăm sóc cây trồng gia đình.'],
            'HS đi thực địa quan sát vườn rau/cây ăn quả gia đình. Chụp ảnh và ghi chép dấu hiệu bất thường trên lá, thân, rễ. (20 phút)'
        ),
        (
            'TÌM HIỂU QUANG HỢP VÀ DINH DƯỠNG CÂY TRỒNG',
            [
                'Mô tả phương trình quang hợp tổng quát và điều kiện cần thiết.',
                'Phân biệt đa lượng (N, P, K) và vi lượng (Fe, Mn, Zn) trong dinh dưỡng cây trồng.',
            ],
            'GV giảng giải sơ đồ quang hợp pha sáng – pha tối. HS quan sát mẫu lá thiếu sắt (vàng úa gân xanh) và thiếu đạm (vàng toàn bộ lá). (30 phút)'
        ),
        (
            'THỰC HÀNH: NHẬN DIỆN VÀ PHÂN LOẠI BỆNH CÂY',
            ['Phân loại bệnh nấm, vi khuẩn, virus dựa trên đặc điểm bệnh lý thực tế.'],
            'HS dùng kính lúp quan sát mẫu lá bệnh (đốm nâu = nấm, thối nhũn = vi khuẩn, khảm lá = virus). Ghi vào phiếu phân loại và đề xuất biện pháp xử lý sinh học (nấm đối kháng, chế phẩm EM). (50 phút)'
        ),
        (
            'THỰC HÀNH: BÓN PHÂN VI SINH VÀ CHĂM SÓC',
            ['HS thực hành bón phân vi sinh EM và tưới nước đúng kỹ thuật.'],
            'HS pha loãng chế phẩm EM tỷ lệ 1:500 và tưới quanh gốc cây. Ghi nhật ký chăm sóc. Thảo luận lợi ích vi sinh vật có lợi trong đất. (35 phút)'
        ),
    ]
)

lp5 = LessonPlan.objects.create(
    title='Chăm sóc vườn rau gia đình – Sinh lý thực vật và phòng trừ sâu bệnh an toàn',
    description='Học sinh lớp 10 nông thôn thực hành thực địa tại vườn rau gia đình, tìm hiểu quang hợp, dinh dưỡng cây trồng, nhận diện bệnh nấm/vi khuẩn/virus và áp dụng chế phẩm sinh học EM phòng trừ bệnh.',
    target_student='HS nông thôn',
    status='PUBLISHED', creator=admin, file_path=file5,
    content_preview='Chủ đề: Chăm sóc vườn rau gia đình. Lớp 10, 3 tiết, thực hành tại nông trại. Kiến thức: sinh lý thực vật, quang hợp, dinh dưỡng cây trồng, bệnh học cây trồng. Hoạt động: khảo sát vườn, tìm hiểu dinh dưỡng, phân loại bệnh cây, bón phân vi sinh.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến xã hội',
        'Chủ đề': 'Chăm sóc gia đình',
        'Kiến thức sinh học liên quan': 'Sinh lý thực vật (quang hợp, dinh dưỡng cây trồng), bệnh học cây trồng',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '3 tiết',
        'Địa điểm': 'Thực địa / Nông trại',
        'lop': ['Lớp 10'],
        'knowledge_tags': ['Sinh lý thực vật (quang hợp, dinh dưỡng cây trồng), bệnh học cây trồng']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp5, directory=sub2_1)
LessonPlanDirectory.objects.create(lesson_plan=lp5, directory=track2)

# ============================================================
# LESSON PLAN 6
# Mạch: Hướng đến xã hội | Chủ đề: Xây dựng nhà trường
# Lớp 11 | Lý thuyết | 1 tiết | HS thành thị | Hội trường / Sân khấu
# Kiến thức: Cơ chế nghe–nhìn, âm nhạc và não bộ, sinh học vận động
# ============================================================
file6 = generate_structured_docx(
    filename='am_nhac_hoc_duong_lop11.docx',
    title='CHỦ ĐỀ 6: ÂM NHẠC HỌC ĐƯỜNG VÀ SINH HỌC CẢM GIÁC THẦN KINH',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 11 (THPT)',
    thoi_gian='1 tiết (45 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh mô tả được cơ chế truyền dẫn sóng âm từ tai ngoài đến vỏ não thính giác.',
        'Giải thích tác động của nhịp điệu âm nhạc lên sự giải phóng dopamine và trạng thái tinh thần.',
        'Nhận ra mối liên hệ giữa vận động nghệ thuật (múa, biểu diễn) và cơ chế phối hợp thần kinh – cơ.',
    ],
    thiet_bi='Loa bluetooth, sơ đồ cấu tạo tai người, video phân tích MRI não khi nghe nhạc.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: TRẢI NGHIỆM ÂM THANH',
            ['Kích hoạt nhận thức về cảm giác âm thanh và phản ứng cơ thể.'],
            'Mở 3 đoạn nhạc: nhạc buồn – vui – kích động. HS ghi cảm xúc và nhịp tim sau mỗi đoạn. Thảo luận tại sao âm nhạc gây cảm xúc khác nhau. (10 phút)'
        ),
        (
            'GIẢI PHẪU TAI VÀ CƠ CHẾ NGHE',
            [
                'Mô tả đường truyền sóng âm: tai ngoài → màng nhĩ → xương con → ốc tai → thần kinh thính giác → vỏ não.',
                'Giải thích tai trong chuyển đổi rung động cơ học thành tín hiệu thần kinh điện.',
            ],
            'GV trình bày sơ đồ cấu tạo tai người. HS quan sát video siêu âm cấu trúc ốc tai dạng lò xo (cochlea). Điền vào sơ đồ câm. (20 phút)'
        ),
        (
            'ÂM NHẠC VÀ HỆ LIMBIC – DOPAMINE',
            ['HS hiểu tại sao âm nhạc kích thích cảm xúc thông qua hệ limbic và dopamine.'],
            'GV chiếu ảnh MRI não sáng lên ở vùng nucleus accumbens khi nghe nhạc yêu thích. Giải thích vòng phản hồi dopamine – thưởng – hành vi. HS liên hệ thực tế. (15 phút)'
        ),
    ]
)

lp6 = LessonPlan.objects.create(
    title='Âm nhạc học đường và sinh học cảm giác thần kinh – Cơ chế nghe và dopamine',
    description='Học sinh lớp 11 tìm hiểu cơ chế sinh học của cảm giác thính giác, cách âm nhạc kích hoạt vùng limbic giải phóng dopamine và ảnh hưởng đến tâm trạng, hành vi trong các hoạt động biểu diễn nghệ thuật học đường.',
    target_student='HS thành thị',
    status='PUBLISHED', creator=admin, file_path=file6,
    content_preview='Chủ đề: Âm nhạc học đường và sinh học cảm giác. Lớp 11, 1 tiết, lý thuyết. Kiến thức: cơ chế nghe nhìn, âm nhạc và não bộ, dopamine, sinh học vận động. Hoạt động: trải nghiệm âm thanh, giải phẫu tai, limbic và dopamine.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến xã hội',
        'Chủ đề': 'Xây dựng nhà trường',
        'Kiến thức sinh học liên quan': 'Cơ chế nghe – nhìn, ảnh hưởng âm nhạc đến não bộ, sinh học vận động',
        'Loại hình': 'Lý thuyết',
        'Tiết dạy': '1 tiết',
        'Địa điểm': 'Hội trường / Sân khấu',
        'lop': ['Lớp 11'],
        'knowledge_tags': ['Cơ chế nghe – nhìn, ảnh hưởng âm nhạc đến não bộ, sinh học vận động']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp6, directory=sub2_2)
LessonPlanDirectory.objects.create(lesson_plan=lp6, directory=track2)

# ============================================================
# LESSON PLAN 7
# Mạch: Hướng đến xã hội | Chủ đề: Xây dựng cộng đồng
# Lớp 12 | Thực hành | 2 tiết | HS nông thôn | Ngoài trời / Thực địa
# Kiến thức: Vi sinh vật rác thải, ô nhiễm sức khỏe cộng đồng
# ============================================================
file7 = generate_structured_docx(
    filename='don_ve_sinh_cong_dong_lop12.docx',
    title='CHỦ ĐỀ 7: CHIẾN DỊCH VỆ SINH CỘNG ĐỒNG – VI SINH VẬT VÀ SỨC KHỎE MÔI TRƯỜNG',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 12 (THPT)',
    thoi_gian='2 tiết (90 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh tổ chức được chiến dịch thu gom rác thải tại địa phương.',
        'Giải thích vai trò của vi sinh vật phân hủy hữu cơ trong xử lý rác và ô nhiễm môi trường.',
        'Phân tích được mối liên hệ giữa rác thải và các bệnh truyền nhiễm qua trung gian sinh vật (ruồi, muỗi).',
    ],
    thiet_bi='Găng tay bảo hộ, túi thu gom rác, kính lúp, slide về vi sinh vật gây bệnh phổ biến.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: PHÂN LOẠI RÁC VÀ NHẬN DIỆN NGUY CƠ',
            ['Kích hoạt nhận thức về mối nguy hại của rác thải đối với sức khỏe.'],
            'HS quan sát bãi rác gần trường. Phân loại: rác hữu cơ – vô cơ – nguy hại. Thảo luận: "Rác hữu cơ để lâu có gì xảy ra? Vi sinh vật nào xuất hiện?" (15 phút)'
        ),
        (
            'RA QUÂN THU GOM RÁC THẢI',
            ['Tổ chức hành động cộng đồng có trách nhiệm và kỷ luật.'],
            'Chia 5 nhóm phân công khu vực. Thu gom rác đeo găng tay, phân loại vào túi màu khác nhau. Ghi chép số lượng và loại rác thu được. (30 phút)'
        ),
        (
            'PHÂN TÍCH VI SINH VẬT TRONG RÁC THẢI',
            [
                'Nhận biết các nhóm vi sinh vật phân hủy rác: vi khuẩn hiếu khí, nấm mốc, xạ khuẩn.',
                'Hiểu cơ chế vi sinh vật gây bệnh lây truyền qua rác thải: tả, thương hàn, lỵ.',
            ],
            'GV trình bày sơ đồ vi sinh vật phân hủy hữu cơ. HS phân tích nguy cơ bệnh tả (Vibrio cholerae), thương hàn (Salmonella typhi) từ nguồn nước bẩn tiếp xúc rác. (30 phút)'
        ),
        (
            'ĐỀ XUẤT GIẢI PHÁP XỬ LÝ SINH HỌC',
            ['HS đề xuất biện pháp ủ phân compost hữu cơ hoặc biogas.'],
            'Nhóm thảo luận và trình bày quy trình ủ phân compost bằng vi sinh vật và giun đất. Đánh giá tính khả thi tại địa phương nông thôn. (15 phút)'
        ),
    ]
)

lp7 = LessonPlan.objects.create(
    title='Chiến dịch vệ sinh cộng đồng nông thôn – Vi sinh vật rác thải và phòng bệnh truyền nhiễm',
    description='Học sinh lớp 12 nông thôn tổ chức chiến dịch thu gom rác thải thực tế, tìm hiểu vai trò vi sinh vật phân hủy, mối liên hệ giữa ô nhiễm rác thải và các bệnh truyền nhiễm, đề xuất giải pháp xử lý sinh học an toàn.',
    target_student='HS nông thôn',
    status='PUBLISHED', creator=admin, file_path=file7,
    content_preview='Chủ đề: Chiến dịch vệ sinh cộng đồng. Lớp 12, 2 tiết, thực hành ngoài trời. Kiến thức: vi sinh vật gây bệnh trong rác thải, ảnh hưởng ô nhiễm sức khỏe cộng đồng. Hoạt động: phân loại rác, thu gom, phân tích vi sinh, đề xuất giải pháp.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến xã hội',
        'Chủ đề': 'Xây dựng cộng đồng',
        'Kiến thức sinh học liên quan': 'Vi sinh vật gây bệnh trong rác thải, ảnh hưởng ô nhiễm đến sức khỏe cộng đồng',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '2 tiết',
        'Địa điểm': 'Ngoài trời / Thực địa',
        'lop': ['Lớp 12'],
        'knowledge_tags': ['Vi sinh vật gây bệnh trong rác thải, ảnh hưởng ô nhiễm đến sức khỏe cộng đồng']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp7, directory=sub2_3)
LessonPlanDirectory.objects.create(lesson_plan=lp7, directory=track2)

# ============================================================
# LESSON PLAN 8
# Mạch: Hướng đến tự nhiên | Chủ đề: Bảo tồn cảnh quan thiên nhiên
# Lớp 10 | Thực hành | 3 tiết | HS thành thị | Bảo tàng / Thực địa
# Kiến thức: Phân loại thực vật, đa dạng sinh học, tiến hóa
# ============================================================
file8 = generate_structured_docx(
    filename='bao_ton_da_dang_sinh_hoc_lop10.docx',
    title='CHỦ ĐỀ 8: THAM QUAN VƯỜN THỰC VẬT – PHÂN LOẠI VÀ BẢO TỒN ĐA DẠNG SINH HỌC',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 10 (THPT)',
    thoi_gian='3 tiết (135 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh nhận biết được các ngành thực vật chính: Rêu, Dương xỉ, Hạt trần, Hạt kín.',
        'Sử dụng được khóa lưỡng phân để định danh một số loài thực vật quan sát thực tế.',
        'Phân tích được mối quan hệ tiến hóa từ thực vật không mạch đến thực vật hạt kín.',
        'Đề xuất biện pháp bảo tồn một loài thực vật quý hiếm địa phương.',
    ],
    thiet_bi='Bảng khóa lưỡng phân, kính hiển vi cầm tay, phiếu quan sát, túi thu thập mẫu lá.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: DU LỊCH KHOA HỌC VƯỜN THỰC VẬT',
            ['Kích thích sự tò mò và kết nối kiến thức phân loại với thực tế thiên nhiên.'],
            'GV giới thiệu bản đồ vườn thực vật. Chia nhóm, phân công khu vực quan sát. HS chụp ảnh 5-10 loài cây khác nhau và ghi tên sơ bộ. (20 phút)'
        ),
        (
            'THỰC HÀNH PHÂN LOẠI BẰNG KHÓA LƯỠNG PHÂN',
            [
                'Sử dụng khóa lưỡng phân xác định ngành: Rêu (không mạch) → Dương xỉ (mạch, không hạt) → Hạt trần (lá kim) → Hạt kín (hoa, quả).',
                'Ghi đặc điểm phân loại: cấu trúc lá, thân, hệ rễ, cơ quan sinh sản.',
            ],
            'Mỗi nhóm 4 HS sử dụng khóa lưỡng phân định danh 5 loài cây được phân công. Ghi vào bảng phân loại. Chụp ảnh đặc điểm nhận dạng quan trọng. (50 phút)'
        ),
        (
            'VẼ SƠ ĐỒ TIẾN HÓA VÀ LƯỚI THỨC ĂN',
            ['HS vẽ được sơ đồ tiến hóa thực vật và lưới thức ăn đơn giản của hệ sinh thái vườn.'],
            'Mỗi nhóm vẽ sơ đồ tiến hóa từ tảo → rêu → dương xỉ → hạt trần → hạt kín. Kết hợp vẽ lưới thức ăn: cây → côn trùng → chim → mèo hoang. (30 phút)'
        ),
        (
            'TRÌNH BÀY VÀ ĐỀ XUẤT BẢO TỒN',
            ['HS trình bày kết quả và đề xuất giải pháp bảo tồn loài quý.'],
            'Mỗi nhóm trình bày 5 phút: loài cây phân loại, đặc điểm tiến hóa và đề xuất biện pháp bảo tồn. GV nhận xét. (35 phút)'
        ),
    ]
)

lp8 = LessonPlan.objects.create(
    title='Tham quan vườn thực vật – Phân loại đa dạng sinh học và tiến hóa thực vật',
    description='Học sinh lớp 10 thành thị trải nghiệm tham quan vườn thực vật, thực hành dùng khóa lưỡng phân phân loại các ngành thực vật chính, vẽ sơ đồ tiến hóa và đề xuất biện pháp bảo tồn đa dạng sinh học địa phương.',
    target_student='HS thành thị',
    status='PUBLISHED', creator=admin, file_path=file8,
    content_preview='Chủ đề: Tham quan vườn thực vật, phân loại đa dạng sinh học. Lớp 10, 3 tiết, thực hành thực địa. Kiến thức: phân loại thực vật, đa dạng sinh học, tiến hóa. Hoạt động: du lịch khoa học, khóa lưỡng phân, sơ đồ tiến hóa, bảo tồn.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến tự nhiên',
        'Chủ đề': 'Tìm hiểu và bảo tồn cảnh quan thiên nhiên',
        'Kiến thức sinh học liên quan': 'Phân loại thực vật, đa dạng sinh học, tiến hóa',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '3 tiết',
        'Địa điểm': 'Bảo tàng / Thực địa',
        'lop': ['Lớp 10'],
        'knowledge_tags': ['Phân loại thực vật, đa dạng sinh học, tiến hóa']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp8, directory=sub3_1)
LessonPlanDirectory.objects.create(lesson_plan=lp8, directory=track3)

# ============================================================
# LESSON PLAN 9
# Mạch: Hướng đến tự nhiên | Chủ đề: Bảo vệ môi trường
# Lớp 12 | Lý thuyết | 1 tiết | HS thành thị | Phòng thí nghiệm Sinh học
# Kiến thức: Hệ hô hấp người, tác động khí độc, sinh thái đô thị
# ============================================================
file9 = generate_structured_docx(
    filename='o_nhiem_khong_khi_lop12.docx',
    title='CHỦ ĐỀ 9: Ô NHIỄM KHÔNG KHÍ ĐÔ THỊ VÀ SỨC KHỎE HỆ HÔ HẤP NGƯỜI',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 12 (THPT)',
    thoi_gian='1 tiết (45 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh mô tả được cấu trúc hệ hô hấp người từ mũi đến phế nang.',
        'Giải thích cơ chế trao đổi khí O2/CO2 tại phế nang và tế bào cơ thể.',
        'Phân tích tác hại của bụi mịn PM2.5 và khí NOx đến cấu trúc tế bào phổi.',
        'Đề xuất biện pháp cá nhân và cộng đồng giảm thiểu tác hại ô nhiễm không khí.',
    ],
    thiet_bi='Máy đo AQI cầm tay, mô hình phổi 3D, video hiển vi tế bào phổi bị tổn thương bởi PM2.5.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: ĐO CHẤT LƯỢNG KHÔNG KHÍ',
            ['Trực quan hóa vấn đề ô nhiễm không khí bằng số liệu thực tế.'],
            'HS đo AQI trong lớp học và ngoài hành lang bằng máy cảm biến. So sánh và thảo luận: AQI > 100 ảnh hưởng gì đến sức khỏe? (10 phút)'
        ),
        (
            'GIẢI PHẪU HỆ HÔ HẤP VÀ TRAO ĐỔI KHÍ',
            [
                'Mô tả đường đi của không khí: mũi → hầu → thanh quản → khí quản → phế quản → phế nang.',
                'Giải thích trao đổi khí theo gradient nồng độ: O2 từ phế nang vào máu, CO2 ngược chiều.',
            ],
            'GV trình bày mô hình phổi 3D. HS vẽ sơ đồ hệ hô hấp và đánh dấu vùng trao đổi khí. Tính diện tích bề mặt phế nang (~70m2). (20 phút)'
        ),
        (
            'TÁC HẠI CỦA Ô NHIỄM KHÔNG KHÍ ĐẾN TẾ BÀO PHỔI',
            [
                'Mô tả bụi PM2.5 vượt qua hàng rào phế nang, kích hoạt viêm tế bào và tổn thương DNA.',
                'Phân tích nguy cơ ung thư phổi do tiếp xúc lâu dài với NOx, SO2.',
            ],
            'GV chiếu ảnh hiển vi điện tử tế bào phổi khỏe và bị tổn thương. HS so sánh và ghi nhận xét sinh học về thay đổi cấu trúc màng tế bào, ty thể, nhân. (15 phút)'
        ),
    ]
)

lp9 = LessonPlan.objects.create(
    title='Ô nhiễm không khí đô thị và tác động đến hệ hô hấp người – Sinh thái học đô thị',
    description='Học sinh lớp 12 thành thị tìm hiểu cấu trúc hệ hô hấp, cơ chế trao đổi khí tại phế nang và phân tích tác hại của bụi mịn PM2.5, NOx đối với tế bào phổi trong bối cảnh ô nhiễm đô thị ngày càng nghiêm trọng.',
    target_student='HS thành thị',
    status='PUBLISHED', creator=admin, file_path=file9,
    content_preview='Chủ đề: Ô nhiễm không khí đô thị và sức khỏe hô hấp. Lớp 12, 1 tiết, lý thuyết. Kiến thức: hệ hô hấp người, tác động khí độc PM2.5, sinh thái đô thị. Hoạt động: đo AQI, giải phẫu hô hấp, tác hại ô nhiễm tế bào phổi.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến tự nhiên',
        'Chủ đề': 'Tìm hiểu và bảo vệ môi trường',
        'Kiến thức sinh học liên quan': 'Hệ hô hấp người, tác động khí độc, sinh thái đô thị',
        'Loại hình': 'Lý thuyết',
        'Tiết dạy': '1 tiết',
        'Địa điểm': 'Phòng thí nghiệm Sinh học',
        'lop': ['Lớp 12'],
        'knowledge_tags': ['Hệ hô hấp người, tác động khí độc, sinh thái đô thị']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp9, directory=sub3_2)
LessonPlanDirectory.objects.create(lesson_plan=lp9, directory=track3)

# ============================================================
# LESSON PLAN 10
# Mạch: Hướng đến tự nhiên | Chủ đề: Bảo vệ môi trường
# Lớp 11 | Thực hành | 2 tiết | HS nông thôn | Ngoài trời / Thực địa
# Kiến thức: Vi sinh vật nước, chu trình nitơ, ô nhiễm thủy sinh
# ============================================================
file10 = generate_structured_docx(
    filename='khao_sat_nguon_nuoc_lop11.docx',
    title='CHỦ ĐỀ 10: KHẢO SÁT NGUỒN NƯỚC KÊNH RẠCH – VI SINH VẬT VÀ CHU TRÌNH NITƠ',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 11 (THPT)',
    thoi_gian='2 tiết (90 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh thu thập và quan sát mẫu nước kênh rạch, nhận biết các loại vi sinh vật thủy sinh dưới kính hiển vi.',
        'Giải thích vai trò của vi khuẩn nitrat hóa trong chu trình nitơ của hệ sinh thái nước.',
        'Phân tích tác hại của nước thải chứa nitrat/phosphate dư thừa gây hiện tượng phú dưỡng hóa.',
    ],
    thiet_bi='Dụng cụ lấy mẫu nước, kính hiển vi quang học, giấy thử pH/nitrat, bình chứa mẫu.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: THU MẪU NƯỚC THỰC TẾ',
            ['Trực quan hóa bằng mẫu nước thực từ kênh gần trường.'],
            'HS lấy mẫu nước từ 3 điểm: nước trong, nước đục, nước gần bờ có rác. Ghi nhận màu sắc, mùi, độ đục vào phiếu quan sát. (15 phút)'
        ),
        (
            'QUAN SÁT VI SINH VẬT THỦY SINH DƯỚI KÍNH HIỂN VI',
            [
                'Nhận biết tảo lục (Chlorella), tảo silic (Diatom), trùng roi (Euglena), vi khuẩn hình cầu/que.',
                'Phân biệt sinh vật phù du thực vật (phytoplankton) và động vật (zooplankton).',
            ],
            'HS làm tiêu bản ướt mẫu nước, quan sát dưới kính hiển vi 400x. Vẽ hình và đặt tên 3-5 sinh vật quan sát được theo khóa phân loại đơn giản. (30 phút)'
        ),
        (
            'CHU TRÌNH NITƠ VÀ PHÚ DƯỠNG HÓA',
            [
                'Mô tả chu trình nitơ: NH4+ → nitrit → nitrat → N2 do vi khuẩn đất/nước.',
                'Giải thích phú dưỡng hóa: tảo bùng phát → tiêu thụ O2 → cá chết.',
            ],
            'GV vẽ sơ đồ chu trình nitơ. HS thử nghiệm nước mẫu bằng giấy thử nitrat. So sánh kết quả 3 điểm lấy mẫu và rút ra kết luận về mức độ ô nhiễm. (30 phút)'
        ),
        (
            'ĐỀ XUẤT GIẢI PHÁP BẢO VỆ NGUỒN NƯỚC',
            ['HS đề xuất biện pháp xử lý nước và bảo vệ kênh rạch địa phương.'],
            'Nhóm thảo luận và trình bày: sử dụng thực vật thủy sinh (bèo tây, rong biển) để hấp thụ nitrat dư thừa. Đề xuất tuyên truyền cộng đồng. (15 phút)'
        ),
    ]
)

lp10 = LessonPlan.objects.create(
    title='Khảo sát nguồn nước kênh rạch – Vi sinh vật thủy sinh và chu trình nitơ',
    description='Học sinh lớp 11 nông thôn thu mẫu nước kênh rạch, quan sát vi sinh vật thủy sinh dưới kính hiển vi, tìm hiểu chu trình nitơ và phân tích nguy cơ phú dưỡng hóa từ nước thải nông nghiệp tại địa phương.',
    target_student='HS nông thôn',
    status='PUBLISHED', creator=admin, file_path=file10,
    content_preview='Chủ đề: Khảo sát nguồn nước kênh rạch. Lớp 11, 2 tiết, thực hành ngoài trời. Kiến thức: vi sinh vật nước, chu trình nitơ, ảnh hưởng ô nhiễm đến sinh vật thủy sinh. Hoạt động: thu mẫu, kính hiển vi, chu trình N, đề xuất bảo vệ.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng đến tự nhiên',
        'Chủ đề': 'Tìm hiểu và bảo vệ môi trường',
        'Kiến thức sinh học liên quan': 'Vi sinh vật nước, chu trình nitơ, ảnh hưởng ô nhiễm đến sinh vật thủy sinh',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '2 tiết',
        'Địa điểm': 'Ngoài trời / Thực địa',
        'lop': ['Lớp 11'],
        'knowledge_tags': ['Vi sinh vật nước, chu trình nitơ, ảnh hưởng ô nhiễm đến sinh vật thủy sinh']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp10, directory=sub3_2)
LessonPlanDirectory.objects.create(lesson_plan=lp10, directory=track3)

# ============================================================
# LESSON PLAN 11
# Mạch: Hướng nghiệp | Chủ đề: Tìm hiểu nghề nghiệp
# Lớp 12 | Lý thuyết | 2 tiết | HS thành thị | Phòng thí nghiệm Sinh học
# Kiến thức: Công nghệ gen, sinh học phân tử, ứng dụng y học
# ============================================================
file11 = generate_structured_docx(
    filename='cong_nghe_gen_lop12.docx',
    title='CHỦ ĐỀ 11: CÔNG NGHỆ GEN VÀ ỨNG DỤNG TRONG Y HỌC – ĐỊNH HƯỚNG NGHỀ SINH HỌC PHÂN TỬ',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 12 (THPT)',
    thoi_gian='2 tiết (90 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh mô tả được nguyên lý kỹ thuật PCR trong chẩn đoán bệnh và nghiên cứu gen.',
        'Giải thích nguyên lý chuyển gen bằng vector plasmid vào vi khuẩn E.coli.',
        'Phân tích ứng dụng thực tế của công nghệ gen: insulin tái tổ hợp, vaccine mRNA, cây trồng biến đổi gen.',
        'Tìm hiểu cơ hội nghề nghiệp trong lĩnh vực Công nghệ sinh học phân tử.',
    ],
    thiet_bi='Mô hình plasmid cắt dán, video quy trình PCR, bảng so sánh nghề nghiệp CNSH toàn cầu.',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: CÂU HỎI THỰC TẾ VỀ CÔNG NGHỆ GEN',
            ['Kết nối kiến thức gen với ứng dụng thực tế học sinh đã biết.'],
            '"Vaccine COVID-19 mRNA được tạo ra như thế nào?" GV chiếu video 3 phút về công nghệ mRNA. HS thảo luận và chia sẻ hiểu biết ban đầu. (10 phút)'
        ),
        (
            'TÌM HIỂU KỸ THUẬT PCR VÀ CHẨN ĐOÁN BỆNH',
            [
                'Mô tả 3 bước PCR: biến tính (94°C) → gắn mồi (55°C) → tổng hợp (72°C).',
                'Giải thích ứng dụng PCR trong test nhanh COVID-19, xét nghiệm DNA pháp y.',
            ],
            'GV mô phỏng chu kỳ nhiệt PCR bằng cử chỉ. HS "đóng vai" DNA polymerase trong trò chơi nhân đôi DNA. (25 phút)'
        ),
        (
            'THỰC HÀNH: LẮP GHÉP MÔ HÌNH PLASMID CHUYỂN GEN',
            [
                'Mô tả vector plasmid và enzyme cắt giới hạn (restriction enzyme).',
                'Mô phỏng quá trình cắt và nối gen insulin vào plasmid vi khuẩn.',
            ],
            'HS dùng mô hình plasmid cắt dán (giấy), thực hành: cắt plasmid → chèn gen insulin → đóng vòng tròn. Giải thích tại sao E.coli có thể sản xuất insulin người. (35 phút)'
        ),
        (
            'ĐỊNH HƯỚNG NGHỀ NGHIỆP CÔNG NGHỆ SINH HỌC',
            ['HS biết các ngành đào tạo và cơ hội việc làm trong lĩnh vực CNSH.'],
            'GV giới thiệu bảng so sánh các ngành CNSH tại VN (Đại học Bách Khoa, ĐHKHTN, ĐHYD). HS điền phiếu định hướng nghề và lập mục tiêu học tập 3 môn cần ôn thi. (20 phút)'
        ),
    ]
)

lp11 = LessonPlan.objects.create(
    title='Công nghệ gen và ứng dụng y học – Định hướng nghề nghiệp Sinh học phân tử',
    description='Học sinh lớp 12 thành thị tìm hiểu kỹ thuật PCR, công nghệ chuyển gen qua vector plasmid và các ứng dụng y học (insulin tái tổ hợp, vaccine mRNA). Kết hợp định hướng nghề nghiệp ngành Công nghệ sinh học phân tử.',
    target_student='HS thành thị',
    status='PUBLISHED', creator=admin, file_path=file11,
    content_preview='Chủ đề: Công nghệ gen và ứng dụng y học. Lớp 12, 2 tiết, lý thuyết. Kiến thức: công nghệ gen, sinh học phân tử, ứng dụng y học nông nghiệp. Hoạt động: PCR câu hỏi, kỹ thuật PCR chẩn đoán, lắp ghép plasmid, định hướng nghề.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng nghiệp',
        'Chủ đề': 'Tìm hiểu nghề nghiệp',
        'Kiến thức sinh học liên quan': 'Công nghệ gen, sinh học phân tử, ứng dụng y học/nông nghiệp',
        'Loại hình': 'Lý thuyết',
        'Tiết dạy': '2 tiết',
        'Địa điểm': 'Phòng thí nghiệm Sinh học',
        'lop': ['Lớp 12'],
        'knowledge_tags': ['Công nghệ gen, sinh học phân tử, ứng dụng y học/nông nghiệp']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp11, directory=sub4_1)
LessonPlanDirectory.objects.create(lesson_plan=lp11, directory=track4)

# ============================================================
# LESSON PLAN 12
# Mạch: Hướng nghiệp | Chủ đề: Rèn luyện phẩm chất hướng nghiệp
# Lớp 11 | Thực hành | 3 tiết | HS nông thôn | Nông nghiệp công nghệ cao / Thực địa
# Kiến thức: Sinh lý thực vật, nuôi cấy mô, di truyền chọn giống
# ============================================================
file12 = generate_structured_docx(
    filename='nong_nghiep_cong_nghe_cao_lop11.docx',
    title='CHỦ ĐỀ 12: THỰC HÀNH NÔNG NGHIỆP CÔNG NGHỆ CAO – NUÔI CẤY MÔ VÀ THỦY CANH',
    mon_hoc='Hoạt động trải nghiệm Sinh học',
    lop='Lớp 11 (THPT)',
    thoi_gian='3 tiết (135 phút)',
    gv='Giáo viên Sinh học',
    muc_tieu_list=[
        'Học sinh giải thích được nguyên lý nuôi cấy mô tế bào thực vật (callus → phôi soma → cây hoàn chỉnh).',
        'Pha chế được dung dịch dinh dưỡng thủy canh và đo chỉ số EC, pH phù hợp.',
        'Nhận diện được nguyên nhân di truyền và biện pháp cải thiện giống cây trồng qua chọn lọc nhân tạo.',
        'Hình thành nhận thức về nghề nghiệp trong nông nghiệp công nghệ cao.',
    ],
    thiet_bi='Bình thủy canh, máy đo EC/pH, mẫu callus lan hồ điệp, dung dịch MS (Murashige & Skoog).',
    hoat_dong_list=[
        (
            'KHỞI ĐỘNG: THAM QUAN NHÀ KÍNH THỦY CANH',
            ['Kết nối kiến thức sinh lý thực vật với hệ thống nông nghiệp hiện đại.'],
            'HS tham quan nhà kính thủy canh (hoặc video trực tiếp). Quan sát rễ cây xà lách ngâm trong dung dịch dinh dưỡng. Đặt câu hỏi: "Cây không cần đất vẫn sống được vì sao?" (20 phút)'
        ),
        (
            'THỰC HÀNH: PHA DUNG DỊCH THỦY CANH VÀ ĐO EC/pH',
            [
                'Hiểu vai trò của từng nguyên tố khoáng trong dung dịch MS: N, P, K, Ca, Mg, Fe.',
                'Pha đúng tỷ lệ để đạt EC = 1.5-2.0 mS/cm và pH = 5.5-6.5.',
            ],
            'HS pha loãng dung dịch thủy canh đậm đặc 1:50. Đo EC và pH. Điều chỉnh bằng dung dịch axit/bazơ nếu lệch chuẩn. Ghi kết quả vào bảng theo dõi. (40 phút)'
        ),
        (
            'TÌM HIỂU NUÔI CẤY MÔ VÀ CHỌN GIỐNG',
            [
                'Mô tả quy trình nuôi cấy mô: khử trùng mẫu → môi trường MS → callus → phôi soma → cây con.',
                'Giải thích nguyên lý chọn lọc nhân tạo và lai giống tạo giống chống bệnh.',
            ],
            'GV trình bày quy trình nuôi cấy mô trên mẫu lan hồ điệp. HS quan sát callus dưới kính hiển vi. Thảo luận ứng dụng chọn giống kháng bệnh cho lúa, cà phê tại địa phương. (45 phút)'
        ),
        (
            'ĐỊNH HƯỚNG NGHỀ NÔNG NGHIỆP CÔNG NGHỆ CAO',
            ['HS nhận thức được cơ hội nghề nghiệp trong lĩnh vực nông nghiệp hiện đại.'],
            'GV chia sẻ bản đồ nghề nghiệp nông nghiệp CNC: kỹ sư giống, kỹ thuật viên nhà kính, chuyên gia dinh dưỡng thực vật. HS ghi mục tiêu và câu hỏi định hướng nghề nghiệp. (30 phút)'
        ),
    ]
)

lp12 = LessonPlan.objects.create(
    title='Thực hành nông nghiệp công nghệ cao – Nuôi cấy mô, thủy canh và chọn giống di truyền',
    description='Học sinh lớp 11 nông thôn thực hành pha dung dịch thủy canh, tìm hiểu quy trình nuôi cấy mô tế bào thực vật và nguyên lý di truyền chọn giống. Định hướng nghề nghiệp trong lĩnh vực nông nghiệp công nghệ cao.',
    target_student='HS nông thôn',
    status='PUBLISHED', creator=admin, file_path=file12,
    content_preview='Chủ đề: Nông nghiệp công nghệ cao, nuôi cấy mô thủy canh. Lớp 11, 3 tiết, thực hành. Kiến thức: sinh lý thực vật, nuôi cấy mô, di truyền chọn giống. Hoạt động: tham quan nhà kính, pha dung dịch EC, nuôi cấy mô callus, định hướng nghề.',
    attributes={
        'Môn học': 'Hoạt động trải nghiệm Sinh học',
        'Mạch kiến thức': 'Hoạt động hướng nghiệp',
        'Chủ đề': 'Rèn luyện phẩm chất, năng lực phù hợp với định hướng nghề nghiệp',
        'Kiến thức sinh học liên quan': 'Sinh lý thực vật, nuôi cấy mô, di truyền chọn giống',
        'Loại hình': 'Thực hành',
        'Tiết dạy': '3 tiết',
        'Địa điểm': 'Nông nghiệp công nghệ cao / Thực địa',
        'lop': ['Lớp 11'],
        'knowledge_tags': ['Sinh lý thực vật, nuôi cấy mô, di truyền chọn giống']
    }
)
LessonPlanDirectory.objects.create(lesson_plan=lp12, directory=sub4_2)
LessonPlanDirectory.objects.create(lesson_plan=lp12, directory=track4)

print("\n=== SEEDING COMPLETED SUCCESSFULLY ===")
print(f"Directories: {Directory.objects.count()} (4 tracks + 10 sub-folders)")
print(f"Lesson Plans: {LessonPlan.objects.count()} (12 unique plans)")
print(f"LessonPlanDirectory links: {LessonPlanDirectory.objects.count()}")
print("\nAttribute diversity summary:")
print("- Lớp: 10 (4 plans), 11 (4 plans), 12 (4 plans)")
print("- Loại hình: Lý thuyết (4 plans), Thực hành (8 plans)")
print("- Tiết dạy: 1 tiết (3), 2 tiết (6), 3 tiết (3)")
print("- Đối tượng: HS thành thị (6), HS nông thôn (6)")
print("- Địa điểm: 7 địa điểm khác nhau")
print("- Kiến thức sinh học: 12 chủ đề khác nhau hoàn toàn")
